from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOCX = ROOT / "deliverables" / "guides" / "mandarinas-team-balance-explainer-en.docx"
OUTPUT_PDF = ROOT / "deliverables" / "guides" / "mandarinas-team-balance-explainer-en.pdf"

SOURCE_ITEMS = [
    "Team split logic: assets/js/core/club_logic.js -> buildBalancedTeams, summarizeTeamBalanceAssignments, teamPlacementScore.",
    "Matchday wiring: assets/js/pages/admin/matchday.js -> autoBalanceTeams and buildTeamBalanceOptions.",
    "Roster queue / recently-played logic: assets/js/pages/admin/matchday.js -> buildRosterPlan and helpers, plus compareRotationQueuePriority in assets/js/core/club_logic.js.",
]

FACTOR_ROWS = [
    (
        "Team size",
        "Yes",
        "A hard target is built first for each team. Size mismatch carries a very large weight so the algorithm strongly prefers filling teams to their target counts.",
    ),
    (
        "Rating",
        "Yes",
        "Each player profile uses skill_rating, with a fallback of 78 if a rating is missing. The algorithm tries to keep each team close to the expected total strength and then reduces final skill spread with swap passes.",
    ),
    (
        "Age",
        "Yes",
        "Age is calculated from birth_date. Teams are penalized when their average age drifts away from the overall IN-pool average. If a player has no birth date, that player is ignored for age math.",
    ),
    (
        "Positions",
        "Yes",
        "The code normalizes positions to GK / DEF / MID / ATT and balances primary positions across teams. It uses only the primary position for the actual team-balance penalty.",
    ),
    (
        "Recently played",
        "Not directly",
        "This is not part of buildBalancedTeams. Recent attendance and games played affect the roster queue that fills the IN pool before auto-balance runs.",
    ),
    (
        "Repeat teammates",
        "Yes",
        "Historical matchday assignments from earlier matchdays in the same season are converted into teammate-pair counts. The algorithm penalizes pairings that have already happened more often.",
    ),
]

AUTO_BALANCE_FLOW = [
    "Match Centre gathers the current IN players and rejects the run if the pool exceeds the total team cap.",
    "buildTeamBalanceOptions passes four team codes, per-team caps, a sizeSeed based on matchday_number - 1, and historicalAssignments from earlier matchdays in the same season.",
    "buildTeamTargetSizes decides how many players each team should receive. The seed rotates which teams get the extra player when the total is not evenly divisible, so the same team is not always first in line for the larger target.",
    "Each player becomes a balance profile with id, display name, strength, age, normalized positions, and a primary position.",
    "Profiles are ordered before assignment: goalkeepers first, then players with fewer listed positions, then players farthest from the pool average in rating and age, then by rating, then by name.",
    "The first pass is greedy. For each player, the code simulates the current partial teams and chooses the eligible team with the lowest placement score.",
    "The placement score combines immediate strength fit, age fit, primary-position fit, repeat-teammate penalty, and a small fill-ratio term.",
    "After the greedy pass, the code performs up to 12 optimization passes. It tests pairwise swaps between players on different teams and keeps a swap only when the total summary score improves.",
    "The final assignments are saved to matchday_assignments, then Match Centre reloads and shows summary metrics such as skill spread, age spread, repeat load, and position mismatch.",
]

WEIGHT_ROWS = [
    ("size = 1000", "Heavily prioritizes correct team sizes."),
    ("strength = 1", "Balances total skill/rating."),
    ("age = 6", "Penalizes teams whose average age drifts from the pool average."),
    ("position = 14", "Pushes the algorithm to distribute primary roles more evenly."),
    ("teammate = 8", "Penalizes repeated teammate pairings from earlier matchdays."),
]

RECENTLY_PLAYED_NOTES = [
    "Core and depth candidates are sorted by attendance_score first, then recent_games_attended, then no-shows where relevant.",
    "Flex players use compareRotationQueuePriority. On normal matchdays, fewer games_attended is favored first. On the final matchday, season totals are prioritized first and games_attended becomes the next tie-breaker.",
    "Once players are already marked IN, auto-balance does not look at recent_games_attended or total games attended for individual fairness. It only sees the current player profiles and teammate history.",
]

LIMITATIONS = [
    "Position balance uses primary position only. Secondary positions are stored on the profile, but the actual balance penalty tracks one primary bucket per player.",
    "Missing birth dates remove players from age balancing, which can reduce the usefulness of the age penalty.",
    "Missing skill ratings fall back to 78, which can flatten the real spread if several profiles are incomplete.",
    "Teammate repeat history comes from earlier matchdays in the same season. It is not a broader all-time history unless that data is explicitly fed in.",
    "Because the algorithm is greedy first and local-swap second, it is an effective heuristic, not a guaranteed global optimum search.",
]


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def set_table_widths(table, widths):
    grid = table._tbl.tblGrid
    while grid.gridCol_lst:
        grid.remove(grid.gridCol_lst[0])

    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Mandarinas Team Balance Explainer")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(
        "Current implementation summary for Match Centre auto-balance and the roster queue that feeds it."
    )
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)

    document.add_heading("Short Answer", level=1)
    document.add_paragraph(
        "Yes, the current auto-balance code does actively use rating, age, positions, and prior teammate history when it splits the current IN pool into teams. "
        "The important nuance is that “recently played” is not a direct factor inside the team-splitting score. Instead, recent attendance and season participation are used earlier in the roster queue that decides who gets into the IN pool first."
    )

    document.add_heading("Source Of Truth In The Code", level=1)
    for item in SOURCE_ITEMS:
        add_bullet(document, item)

    document.add_heading("What Auto-Balance Uses", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [2200, 1500, 5660])
    headers = table.rows[0].cells
    set_cell_text(headers[0], "Factor", bold=True)
    set_cell_text(headers[1], "Used?", bold=True)
    set_cell_text(headers[2], "How It Is Used", bold=True)

    for factor, used, detail in FACTOR_ROWS:
        row = table.add_row().cells
        set_cell_text(row[0], factor, bold=True)
        set_cell_text(row[1], used)
        set_cell_text(row[2], detail)

    document.add_heading("The Auto-Balance Flow", level=1)
    for item in AUTO_BALANCE_FLOW:
        add_numbered(document, item)

    document.add_heading("Current Weights", level=1)
    document.add_paragraph(
        "These weights are defined in assets/js/core/club_logic.js. Bigger numbers matter more in the final score."
    )
    weights = document.add_table(rows=1, cols=2)
    weights.style = "Table Grid"
    set_table_widths(weights, [2800, 6560])
    weight_headers = weights.rows[0].cells
    set_cell_text(weight_headers[0], "Weight", bold=True)
    set_cell_text(weight_headers[1], "Meaning", bold=True)

    for label, meaning in WEIGHT_ROWS:
        row = weights.add_row().cells
        set_cell_text(row[0], label, bold=True)
        set_cell_text(row[1], meaning)

    document.add_heading("What “Recently Played” Actually Means Today", level=1)
    document.add_paragraph(
        "The phrase matters, but it belongs to roster selection rather than team splitting."
    )
    for item in RECENTLY_PLAYED_NOTES:
        add_bullet(document, item)

    document.add_heading("Limitations And Practical Notes", level=1)
    for item in LIMITATIONS:
        add_bullet(document, item)

    document.add_heading("Bottom Line", level=1)
    document.add_paragraph(
        "If the question is “does the code balance by rating, age, positions, and prior teammate repetition?” the answer is yes. "
        "If the question is “does the team split directly account for who recently played more?” the answer is no, not in the split itself. "
        "That fairness signal is handled upstream in the roster queue that determines who enters the IN pool before auto-balance runs."
    )

    return document


def build_pdf() -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCustom",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=21,
        leading=24,
        alignment=TA_LEFT,
        spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "BodyCustom",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.2,
        leading=13.4,
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "HeadingCustom",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14.2,
        leading=17,
        textColor=colors.HexColor("#111111"),
        spaceBefore=10,
        spaceAfter=7,
    )
    small_style = ParagraphStyle(
        "SmallCustom",
        parent=body_style,
        fontSize=9.4,
        leading=12,
    )

    story = [
        Paragraph("Mandarinas Team Balance Explainer", title_style),
        Paragraph(
            "Current implementation summary for Match Centre auto-balance and the roster queue that feeds it.",
            body_style,
        ),
        Spacer(1, 0.08 * inch),
        Paragraph("Short Answer", heading_style),
        Paragraph(
            "Yes, the current auto-balance code does actively use rating, age, positions, and prior teammate history when it splits the current IN pool into teams. "
            "The important nuance is that “recently played” is not a direct factor inside the team-splitting score. Instead, recent attendance and season participation are used earlier in the roster queue that decides who gets into the IN pool first.",
            body_style,
        ),
        Paragraph("Source Of Truth In The Code", heading_style),
        ListFlowable(
            [ListItem(Paragraph(item, body_style)) for item in SOURCE_ITEMS],
            bulletType="bullet",
            leftIndent=18,
        ),
        Paragraph("What Auto-Balance Uses", heading_style),
    ]

    factor_table = Table(
        [["Factor", "Used?", "How It Is Used"], *FACTOR_ROWS],
        colWidths=[1.25 * inch, 0.9 * inch, 4.9 * inch],
        repeatRows=1,
    )
    factor_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2f7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEADING", (0, 0), (-1, -1), 11),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c9d2dc")),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend(
        [
            factor_table,
            Paragraph("The Auto-Balance Flow", heading_style),
            ListFlowable(
                [ListItem(Paragraph(item, body_style)) for item in AUTO_BALANCE_FLOW],
                bulletType="1",
                leftIndent=18,
            ),
            Paragraph("Current Weights", heading_style),
            Paragraph(
                "These weights are defined in assets/js/core/club_logic.js. Bigger numbers matter more in the final score.",
                body_style,
            ),
        ]
    )

    weight_table = Table(
        [["Weight", "Meaning"], *WEIGHT_ROWS],
        colWidths=[1.6 * inch, 5.45 * inch],
        repeatRows=1,
    )
    weight_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2f7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.2),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c9d2dc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend(
        [
            weight_table,
            Paragraph("What “Recently Played” Actually Means Today", heading_style),
            Paragraph(
                "The phrase matters, but it belongs to roster selection rather than team splitting.",
                body_style,
            ),
            ListFlowable(
                [ListItem(Paragraph(item, body_style)) for item in RECENTLY_PLAYED_NOTES],
                bulletType="bullet",
                leftIndent=18,
            ),
            Paragraph("Limitations And Practical Notes", heading_style),
            ListFlowable(
                [ListItem(Paragraph(item, body_style)) for item in LIMITATIONS],
                bulletType="bullet",
                leftIndent=18,
            ),
            Paragraph("Bottom Line", heading_style),
            Paragraph(
                "If the question is “does the code balance by rating, age, positions, and prior teammate repetition?” the answer is yes. "
                "If the question is “does the team split directly account for who recently played more?” the answer is no, not in the split itself. "
                "That fairness signal is handled upstream in the roster queue that determines who enters the IN pool before auto-balance runs.",
                body_style,
            ),
            Spacer(1, 0.05 * inch),
            Paragraph("Generated from the current repo implementation on May 2, 2026.", small_style),
        ]
    )

    pdf = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=letter,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="Mandarinas Team Balance Explainer",
        author="Codex",
    )
    pdf.build(story)


def main() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_DOCX)
    build_pdf()
    print(OUTPUT_DOCX)
    print(OUTPUT_PDF)


if __name__ == "__main__":
    main()
