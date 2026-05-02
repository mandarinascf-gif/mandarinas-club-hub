#!/usr/bin/env python3
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

SKILL_SCRIPTS = Path(
    "/Users/andresaguayo/.codex/plugins/cache/openai-primary-runtime/documents/26.430.10722/skills/documents/scripts"
)
if str(SKILL_SCRIPTS) not in sys.path:
    sys.path.append(str(SKILL_SCRIPTS))

from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa


ROOT = Path("/Users/andresaguayo/mandarinas-club-hub")
OUTPUT_DIR = ROOT / "deliverables" / "guides"
OUTPUT_BASENAME = "mandarinas-rotation-queue-explainer-en"
DOCX_PATH = OUTPUT_DIR / f"{OUTPUT_BASENAME}.docx"
PDF_PATH = OUTPUT_DIR / f"{OUTPUT_BASENAME}.pdf"

TITLE_COLOR = RGBColor(31, 75, 50)
ACCENT_COLOR = RGBColor(196, 90, 26)
MUTED_COLOR = RGBColor(96, 108, 118)
LIGHT_RULE = "D9E2DA"
HEADER_FILL = "EAF2EC"

META_ROWS = [
    ("Audience", "Club admins, captains, and players who want to understand the live queue."),
    ("Scope", "The live rotation queue only. This does not describe the future tier recommendation engine."),
    ("Source of truth", "The current app logic used by Tier Watch, the admin Tiers page, and matchday priority fill."),
    ("Logic reviewed", "May 1, 2026."),
]

WHERE_USED = [
    "Public Tier Watch: shows the live rotation order and the reason each player is where they are.",
    "Admin Tiers page: shows the same live queue alongside the recommendation board.",
    "Matchday priority fill: uses the same queue logic when the app sorts rotation candidates for roster planning.",
]

MODE_STEPS = [
    "The app looks for the next open matchday in the selected season rather than using a generic calendar date.",
    "If that next open matchday is the season's last matchday, the queue switches into final-matchday mode.",
    "The last matchday comes from the season total if it is configured; otherwise the app uses the highest matchday number in the season.",
    "The standings metrics that feed the queue come from completed matchdays only, so unfinished matchdays do not distort the order mid-night.",
]

MODE_SUMMARY_ROWS = [
    (
        "When it activates",
        "The next open matchday is not the last matchday of the season.",
        "The next open matchday is the last matchday of the season.",
    ),
    (
        "Who appears",
        "Eligible rotation-tier players only.",
        "Eligible rotation-tier players only.",
    ),
    (
        "First decision key",
        "Fewest season matchdays attended.",
        "Most season total points.",
    ),
    (
        "Main tiebreak path",
        "After attendance, the queue uses season performance: points, PPG, wins, draws, fewer losses, goals, goalkeeping, clean sheets, then attendance points.",
        "After points, the queue uses season performance: PPG, wins, draws, fewer losses, goals, goalkeeping, clean sheets, then attendance points.",
    ),
    (
        "Later fallbacks",
        "Attendance score, fewer no-shows, fewer late cancels, then player name.",
        "Games attended, attendance score, fewer no-shows, fewer late cancels, then player name.",
    ),
    (
        "Operational meaning",
        "The app favors rotation fairness first, then rewards stronger results inside the tied attendance group.",
        "The app shifts to season results first so the last-night queue reflects who produced the strongest season.",
    ),
]

REGULAR_STEPS = [
    "Include only players whose current season tier is rotation and whose eligibility flag is true.",
    "Sort those players by the fewest season matchdays attended.",
    "If two players are tied on attendance, compare their season performance in this order: total points, points per game, wins, draws, fewer losses, goals, goalkeeping points, clean sheets, then attendance points.",
    "If they are still tied, compare attendance score.",
    "If they are still tied after that, prefer fewer no-shows, then fewer late cancels, then alphabetical name order.",
]

FINAL_STEPS = [
    "Include only players whose current season tier is rotation and whose eligibility flag is true.",
    "Sort first by total points.",
    "If players are tied on points, compare points per game, wins, draws, fewer losses, goals, goalkeeping points, clean sheets, then attendance points.",
    "If players are still tied after the performance checks, use fewer games attended next.",
    "After that, use attendance score, fewer no-shows, fewer late cancels, then alphabetical name order.",
]

EXAMPLES = [
    (
        "1",
        "Regular week: Player A attended 3 matchdays and has 4 points. Player B attended 4 matchdays and has 9 points.",
        "Player A stays ahead because regular mode uses fewer attended matchdays before points.",
    ),
    (
        "2",
        "Final week: the same two players are compared for the last open spot.",
        "Player B jumps ahead because final mode uses season points before games attended.",
    ),
    (
        "3",
        "Two players are tied on the main keys and one has more no-shows or late cancels.",
        "The more reliable player is placed first once the queue reaches those fallback checks.",
    ),
]

PRACTICAL_NOTES = [
    "Core and Sub players are not part of the rotation queue. The queue is rotation-tier only.",
    "Eligibility matters. A rotation player marked ineligible is excluded even if they would otherwise rank near the top.",
    "The queue shown in the app is season-specific. Changing seasons changes the queue context, statistics, and final-matchday calculation.",
    "If a player asks why they moved, the quickest explanation is usually either fewer appearances on a regular week or stronger season results on the final week.",
    "If you want the shortest staff summary: regular weeks are attendance-first; the final week is points-first.",
]


def set_paragraph_border(paragraph, *, bottom: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    if bottom:
        bottom_el = p_bdr.find(qn("w:bottom"))
        if bottom_el is None:
            bottom_el = OxmlElement("w:bottom")
            p_bdr.append(bottom_el)
        bottom_el.set(qn("w:val"), "single")
        bottom_el.set(qn("w:sz"), "8")
        bottom_el.set(qn("w:space"), "4")
        bottom_el.set(qn("w:color"), bottom)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def style_doc(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.4)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(24, 28, 30)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title_style = document.styles["Title"]
    title_style.font.name = "Arial"
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = TITLE_COLOR
    title_style.paragraph_format.space_after = Pt(4)

    subtitle_style = document.styles["Subtitle"]
    subtitle_style.font.name = "Arial"
    subtitle_style.font.size = Pt(11)
    subtitle_style.font.color.rgb = MUTED_COLOR
    subtitle_style.paragraph_format.space_after = Pt(12)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Arial"
    heading1.font.size = Pt(15)
    heading1.font.bold = True
    heading1.font.color.rgb = TITLE_COLOR
    heading1.paragraph_format.space_before = Pt(16)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.keep_with_next = True

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Arial"
    heading2.font.size = Pt(12.5)
    heading2.font.bold = True
    heading2.font.color.rgb = ACCENT_COLOR
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(4)
    heading2.paragraph_format.keep_with_next = True


def add_header_and_footer(document: Document) -> None:
    section = document.sections[0]

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Mandarinas Club Hub | Flex Queue Explainer")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED_COLOR
    set_paragraph_border(p, bottom=LIGHT_RULE)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(2)
    fr = fp.add_run("Page ")
    fr.font.name = "Arial"
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED_COLOR
    add_page_number(fp)


def add_intro(document: Document) -> None:
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Mandarinas Club Hub Flex Queue")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.add_run(
        "Plain-English explanation of how the live club app orders eligible rotation players for the next open spot."
    )

    add_key_value_table(document, META_ROWS)


def add_key_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    section = document.sections[0]
    widths = column_widths_from_weights((1.45, 4.75), section_content_width_dxa(section))
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row()
        left, right = row.cells
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(left, HEADER_FILL)
        p1 = left.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = TITLE_COLOR
        p2 = right.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        r2.font.name = "Arial"
        r2.font.size = Pt(10.5)
    apply_table_geometry(table, widths, table_width_dxa=sum(widths))
    document.add_paragraph("")


def add_body_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def add_number(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def add_mode_summary_table(document: Document) -> None:
    section = document.sections[0]
    widths = column_widths_from_weights((1.8, 2.9, 2.9), section_content_width_dxa(section))
    table = document.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    header = table.add_row().cells
    for cell, text in zip(header, ("Question", "Regular matchday", "Final matchday")):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = TITLE_COLOR

    for row_values in MODE_SUMMARY_ROWS:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = row[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.name = "Arial"
            r.font.size = Pt(10.25)

    apply_table_geometry(table, widths, table_width_dxa=sum(widths))
    document.add_paragraph("")


def add_examples_table(document: Document) -> None:
    section = document.sections[0]
    widths = column_widths_from_weights((1.1, 2.45, 2.45), section_content_width_dxa(section))
    table = document.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    header = table.add_row().cells
    for cell, text in zip(header, ("Example", "Scenario", "Outcome")):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = TITLE_COLOR

    for values in EXAMPLES:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            cell = row[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.name = "Arial"
            r.font.size = Pt(10.25)
            if idx == 0:
                r.bold = True
                r.font.color.rgb = ACCENT_COLOR

    apply_table_geometry(table, widths, table_width_dxa=sum(widths))
    document.add_paragraph("")


def build_document() -> Document:
    document = Document()
    style_doc(document)
    add_header_and_footer(document)
    add_intro(document)

    document.add_heading("What the rotation queue is", level=1)
    add_body_paragraph(
        document,
        "The rotation queue is the live order the club app uses when deciding which eligible rotation player is next for an open spot. In plain terms: it is the app's current answer to, 'Who should be up next in the rotation line?'",
    )
    add_body_paragraph(
        document,
        "The queue is intentionally separate from future tier recommendations. A player can appear high or low in today's queue without that automatically changing their long-term tier recommendation.",
    )

    document.add_heading("Where the queue is used", level=1)
    for item in WHERE_USED:
        add_bullet(document, item)

    document.add_heading("How the app decides which queue mode to use", level=1)
    for item in MODE_STEPS:
        add_number(document, item)

    document.add_heading("Regular vs final matchday", level=1)
    add_mode_summary_table(document)

    document.add_heading("Regular matchday queue logic", level=1)
    add_body_paragraph(
        document,
        "On a normal week, the queue is built to protect rotation fairness first. The first question is not who has the most points. The first question is who has had fewer turns.",
    )
    for item in REGULAR_STEPS:
        add_number(document, item)

    document.add_heading("Final matchday queue logic", level=1)
    add_body_paragraph(
        document,
        "On the season's last matchday, the queue flips the opening priority. The app stops leading with fairness-by-appearances and starts leading with season performance.",
    )
    for item in FINAL_STEPS:
        add_number(document, item)

    document.add_heading("Examples", level=1)
    add_examples_table(document)

    document.add_heading("Practical notes for staff", level=1)
    for item in PRACTICAL_NOTES:
        add_bullet(document, item)

    return document


def pdf_styles():
    palette = {
        "title": colors.HexColor("#1F4B32"),
        "accent": colors.HexColor("#C45A1A"),
        "muted": colors.HexColor("#606C76"),
        "border": colors.HexColor("#D9E2DA"),
        "header_fill": colors.HexColor("#EAF2EC"),
        "text": colors.HexColor("#1A1E1F"),
    }
    styles = getSampleStyleSheet()
    return {
        "palette": palette,
        "title": ParagraphStyle(
            "RQTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=23,
            leading=27,
            textColor=palette["title"],
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "subtitle": ParagraphStyle(
            "RQSubtitle",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            textColor=palette["muted"],
            spaceAfter=12,
        ),
        "heading1": ParagraphStyle(
            "RQHeading1",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=14.5,
            leading=18,
            textColor=palette["title"],
            spaceBefore=12,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "RQBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            textColor=palette["text"],
            spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "RQBullet",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            leftIndent=14,
            firstLineIndent=0,
            bulletIndent=0,
            textColor=palette["text"],
            spaceAfter=4,
        ),
        "table": ParagraphStyle(
            "RQTable",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.4,
            leading=12,
            textColor=palette["text"],
        ),
        "table_header": ParagraphStyle(
            "RQTableHeader",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9.4,
            leading=12,
            textColor=palette["title"],
        ),
        "label": ParagraphStyle(
            "RQLabel",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9.4,
            leading=12,
            textColor=palette["title"],
        ),
    }


def make_bullet_list(items: list[str], style: ParagraphStyle, bullet_type: str = "bullet", start: str = "1"):
    kwargs = {
        "leftIndent": 16,
        "bulletFontName": "Helvetica",
        "bulletFontSize": 10,
        "bulletDedent": 8,
        "spaceBefore": 0,
        "spaceAfter": 2,
    }
    if bullet_type == "bullet":
        kwargs["bulletType"] = "bullet"
        kwargs["bulletChar"] = "•"
    else:
        kwargs["bulletType"] = bullet_type
        kwargs["start"] = start
    return ListFlowable([ListItem(Paragraph(item, style)) for item in items], **kwargs)


def pdf_table(data: list[list[str]], col_widths: list[float], styles_map, first_col_fill=False) -> Table:
    table_data = []
    for row_idx, row in enumerate(data):
        converted = []
        for col_idx, value in enumerate(row):
            style_key = "table_header" if row_idx == 0 else ("label" if first_col_fill and col_idx == 0 else "table")
            converted.append(Paragraph(value, styles_map[style_key]))
        table_data.append(converted)

    table = Table(table_data, colWidths=col_widths, repeatRows=1 if len(table_data) > 1 else 0)
    palette = styles_map["palette"]
    style_commands = [
        ("BOX", (0, 0), (-1, -1), 0.75, palette["border"]),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, palette["border"]),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BACKGROUND", (0, 0), (-1, 0), palette["header_fill"]),
    ]
    if first_col_fill:
        style_commands.append(("BACKGROUND", (0, 1), (0, -1), palette["header_fill"]))
    table.setStyle(TableStyle(style_commands))
    return table


def draw_header_footer(canvas, doc) -> None:
    canvas.saveState()
    width, height = letter
    left = doc.leftMargin
    right = width - doc.rightMargin
    top = height - 0.56 * inch
    canvas.setFont("Helvetica", 8.5)
    canvas.setFillColor(colors.HexColor("#606C76"))
    canvas.drawString(left, top, "Mandarinas Club Hub | Flex Queue Explainer")
    canvas.setStrokeColor(colors.HexColor("#D9E2DA"))
    canvas.setLineWidth(0.7)
    canvas.line(left, top - 6, right, top - 6)
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(right, 0.48 * inch, f"Page {canvas.getPageNumber()}")
    canvas.restoreState()


def build_pdf() -> None:
    styles_map = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.95 * inch,
        rightMargin=0.95 * inch,
        topMargin=1.0 * inch,
        bottomMargin=0.75 * inch,
        title="Mandarinas Club Hub Flex Queue",
        author="OpenAI Codex",
    )
    story = []
    story.append(Paragraph("Mandarinas Club Hub Flex Queue", styles_map["title"]))
    story.append(
        Paragraph(
            "Plain-English explanation of how the live club app orders eligible rotation players for the next open spot.",
            styles_map["subtitle"],
        )
    )

    story.append(
        pdf_table(
            [[label, value] for label, value in META_ROWS],
            [1.55 * inch, 5.95 * inch],
            styles_map,
            first_col_fill=True,
        )
    )
    story.append(Spacer(1, 0.12 * inch))

    story.append(Paragraph("What the rotation queue is", styles_map["heading1"]))
    story.append(
        Paragraph(
            "The rotation queue is the live order the club app uses when deciding which eligible rotation player is next for an open spot. In plain terms: it is the app's current answer to, 'Who should be up next in the rotation line?'",
            styles_map["body"],
        )
    )
    story.append(
        Paragraph(
            "The queue is intentionally separate from future tier recommendations. A player can appear high or low in today's queue without that automatically changing their long-term tier recommendation.",
            styles_map["body"],
        )
    )

    story.append(Paragraph("Where the queue is used", styles_map["heading1"]))
    story.append(make_bullet_list(WHERE_USED, styles_map["bullet"]))

    story.append(Paragraph("How the app decides which queue mode to use", styles_map["heading1"]))
    story.append(make_bullet_list(MODE_STEPS, styles_map["bullet"], bullet_type="1"))

    story.append(Paragraph("Regular vs final matchday", styles_map["heading1"]))
    story.append(
        pdf_table(
            [["Question", "Regular matchday", "Final matchday"], *MODE_SUMMARY_ROWS],
            [1.7 * inch, 2.88 * inch, 2.88 * inch],
            styles_map,
        )
    )
    story.append(Spacer(1, 0.1 * inch))

    story.append(Paragraph("Regular matchday queue logic", styles_map["heading1"]))
    story.append(
        Paragraph(
            "On a normal week, the queue is built to protect rotation fairness first. The first question is not who has the most points. The first question is who has had fewer turns.",
            styles_map["body"],
        )
    )
    story.append(make_bullet_list(REGULAR_STEPS, styles_map["bullet"], bullet_type="1"))

    story.append(Paragraph("Final matchday queue logic", styles_map["heading1"]))
    story.append(
        Paragraph(
            "On the season's last matchday, the queue flips the opening priority. The app stops leading with fairness-by-appearances and starts leading with season performance.",
            styles_map["body"],
        )
    )
    story.append(make_bullet_list(FINAL_STEPS, styles_map["bullet"], bullet_type="1"))

    story.append(Paragraph("Examples", styles_map["heading1"]))
    story.append(
        pdf_table(
            [["Example", "Scenario", "Outcome"], *EXAMPLES],
            [0.8 * inch, 3.45 * inch, 3.25 * inch],
            styles_map,
        )
    )
    story.append(Spacer(1, 0.1 * inch))

    story.append(Paragraph("Practical notes for staff", styles_map["heading1"]))
    story.append(make_bullet_list(PRACTICAL_NOTES, styles_map["bullet"]))

    doc.build(story, onFirstPage=draw_header_footer, onLaterPages=draw_header_footer)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(DOCX_PATH)
    build_pdf()
    print(PDF_PATH)


if __name__ == "__main__":
    main()
